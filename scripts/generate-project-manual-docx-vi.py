from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION
from datetime import date

output_path = r"e:\Study\C#\managerCMN\docs\manual\HUONG_DAN_DU_AN_STAFF_MANAGER_TIENG_VIET.docx"

accented_title = "HƯỚNG DẪN SỬ DỤNG CHỨC NĂNG DỰ ÁN (PROJECT) - BẢN DÀNH CHO STAFF / MANAGER"

doc = Document()

# Set default font
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

for name in ['Title', 'Heading 1', 'Heading 2', 'Heading 3']:
    try:
        style_obj = doc.styles[name]
        style_obj.font.name = 'Times New Roman'
    except Exception:
        pass


def add_text(text, bold=False, size=None, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    if size:
        r.font.size = Pt(size)
    r.font.name = 'Times New Roman'
    return p


def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(18)
    return p


def add_heading(text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = 'Times New Roman'
    return p


def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    return p


def add_number(text):
    p = doc.add_paragraph(style='List Number')
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    return p


def add_placeholder(code, desc):
    p = doc.add_paragraph()
    r = p.add_run(f"[VI TRI CHEN ANH - {code}]\n{desc}")
    r.bold = True
    r.font.name = 'Times New Roman'
    return p


def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, cell_text in enumerate(row):
            cells[i].text = cell_text
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.name = 'Times New Roman'
    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = Cm(width)
    return table


add_title(accented_title)
add_text(f"Hệ thống: managerCMN | Ngày cập nhật: {date.today().strftime('%d/%m/%Y')}", align=WD_PARAGRAPH_ALIGNMENT.CENTER)
add_text("Tài liệu này được viết cho người dùng nội bộ cần nắm luồng thao tác của module Dự án theo đúng phân quyền hiện có trong ứng dụng. Nội dung tập trung vào cách dùng thực tế, kèm vị trí chèn ảnh để bạn tự bổ sung screenshot khi cần.")

add_heading('1. Mục tiêu tài liệu', 1)
add_bullet('Hướng dẫn người dùng hiểu nhanh chức năng Dự án trong managerCMN.')
add_bullet('Tách rõ luồng sử dụng cho Owner, Manager và Staff.')
add_bullet('Ghi rõ các ràng buộc nghiệp vụ để hạn chế thao tác sai.')
add_bullet('Cung cấp mô tả ảnh để bạn tự chèn ảnh chụp màn hình.')
add_placeholder('P00', 'Ảnh bìa hoặc ảnh chụp menu Dự án trong thanh điều hướng bên trái.')

add_heading('2. Vai trò và quyền hạn', 1)
add_table(
    ['Vai trò', 'Quyền chính', 'Ghi chú sử dụng'],
    [
        ['Admin hệ thống', 'Xem tất cả dự án, lưu trữ, xoá vĩnh viễn, quản lý thành viên và task.', 'Có quyền cao nhất trong module Dự án.'],
        ['Project Owner', 'Quản lý dự án, thành viên, task, lưu trữ dự án.', 'Là người chịu trách nhiệm chính của dự án.'],
        ['Project Manager', 'Xem phạm vi công việc được quản lý, theo dõi task, cập nhật tiến độ khi được giao.', 'Giao diện thường giống staff nhưng tập trung vào phần công việc được phân quyền.'],
        ['Project Staff', 'Xem task được giao, mở chi tiết task, cập nhật phần việc theo quyền của mình.', 'Thường là người thực thi công việc hằng ngày.'],
        ['Project Viewer', 'Chỉ xem.', 'Không chỉnh sửa dữ liệu.']
    ],
    col_widths=[4, 7, 6]
)

add_heading('3. Cấu trúc màn hình cần hiểu', 1)
add_bullet('Danh sách dự án: nơi vào module Dự án, lọc trạng thái và mở chi tiết.')
add_bullet('Chi tiết dự án: màn hình chính gồm KPI, tab Công việc, tab Thành viên, panel chi tiết task.')
add_bullet('Tab Công việc: cây task/subtask, timeline/Gantt, thao tác mở chi tiết task.')
add_bullet('Panel chi tiết task: gồm Tổng quan, Subtask & Checklist, Nhật ký.')
add_bullet('Tab Thành viên: thêm thành viên, đổi vai trò, xoá thành viên.')
add_placeholder('P01', 'Ảnh chụp màn hình Danh sách dự án với các thẻ dự án và bộ lọc trạng thái.')

add_heading('4. Hướng dẫn cho Manager', 1)
add_bullet('Manager chủ yếu dùng để theo dõi tiến độ, kiểm tra các task mình quản lý hoặc được giao và phối hợp với Owner.')
add_bullet('Ở màn hình chi tiết dự án, Manager thường vào chế độ xem công việc theo phạm vi được phân quyền, thay vì quản trị toàn bộ cấu hình dự án.')
add_bullet('Nếu được Owner cho phép thao tác ở mức task, Manager có thể mở chi tiết task để xem mô tả, checklist, nhật ký và cập nhật tiến độ theo luồng được cấp quyền.')
add_bullet('Khi làm việc, Manager nên ưu tiên kiểm tra timeline, task quá hạn, task đang review và các subtask đang chờ xử lý.')
add_number('Vào Dự án > chọn dự án cần theo dõi.')
add_number('Mở tab Công việc để xem toàn bộ cây công việc hoặc phần công việc được phép xem.')
add_number('Bấm vào tên task hoặc thanh timeline để mở panel chi tiết task.')
add_number('Đọc kỹ tiến độ, người thực hiện, subtask và checklist trước khi báo cáo trạng thái.')
add_placeholder('P02', 'Ảnh chụp tab Công việc ở chế độ Manager, thể hiện task tree và khu vực timeline.')
add_placeholder('P03', 'Ảnh chụp panel chi tiết task với tab Tổng quan đang mở.')

add_heading('5. Hướng dẫn cho Staff', 1)
add_bullet('Staff dùng module Dự án chủ yếu để nhận việc, xem chi tiết task và cập nhật phần việc được giao.')
add_bullet('Staff không nên thao tác vào phần quản trị thành viên hoặc các nút chỉ dành cho Owner/Admin.')
add_bullet('Khi nhận task, cần đọc mô tả, hạn hoàn thành, checklist và nhật ký trước khi thực hiện.')
add_bullet('Nếu task được giao nhiều người, cần xem phần người thực hiện để biết trách nhiệm của mình nằm ở đâu.')
add_number('Vào Dự án > mở dự án cần làm.')
add_number('Trong tab Công việc, tìm task có tên của bạn trong danh sách người thực hiện hoặc dùng bộ lọc nếu có.')
add_number('Bấm vào task để mở panel chi tiết.')
add_number('Làm theo mô tả ở tab Tổng quan, cập nhật checklist nếu được phép, và theo dõi nhật ký.')
add_number('Nếu task có nhiều subtask, xem từng subtask trước khi báo hoàn thành.')
add_placeholder('P04', 'Ảnh chụp tab Công việc ở góc nhìn Staff, nhấn vào tên task để mở chi tiết.')
add_placeholder('P05', 'Ảnh chụp panel chi tiết task - tab Subtask & Checklist, thể hiện các mục cần làm.')

add_heading('6. Luồng thao tác chi tiết', 1)
add_heading('6.1 Danh sách dự án', 2)
add_number('Vào menu Dự án ở thanh bên trái.')
add_number('Dùng bộ lọc trạng thái để thu hẹp danh sách.')
add_number('Nhìn các chỉ số tóm tắt như % hoàn thành, số task, số thành viên, số task quá hạn.')
add_number('Bấm vào thẻ dự án để vào chi tiết.')
add_placeholder('P06', 'Ảnh chụp trang Danh sách dự án với bộ lọc trạng thái và các thẻ dự án.')

add_heading('6.2 Tạo và sửa dự án', 2)
add_number('Chỉ Owner hoặc Admin mới nên tạo/sửa thông tin dự án.')
add_number('Nhập tên, mô tả, ngày bắt đầu và ngày kết thúc.')
add_number('Kiểm tra lại trạng thái trước khi lưu.')
add_number('Sau khi lưu, quay lại chi tiết để kiểm tra dữ liệu hiển thị đúng.')
add_placeholder('P07', 'Ảnh chụp form Tạo dự án mới.')
add_placeholder('P08', 'Ảnh chụp form Sửa dự án.')

add_heading('6.3 Quản lý thành viên', 2)
add_number('Mở tab Thành viên.')
add_number('Bấm Thêm thành viên nếu có quyền quản lý.')
add_number('Tìm kiếm theo tên, mã nhân viên hoặc phòng ban.')
add_number('Chọn vai trò phù hợp: Manager, Staff hoặc Viewer.')
add_number('Dùng menu ba chấm để đổi vai trò hoặc xoá khỏi dự án.')
add_placeholder('P09', 'Ảnh chụp tab Thành viên của dự án.')
add_placeholder('P10', 'Ảnh chụp popup Thêm thành viên, có ô tìm kiếm và bộ lọc phòng ban.')

add_heading('6.4 Quản lý task', 2)
add_number('Mở tab Công việc.')
add_number('Dùng bộ lọc trạng thái và người thực hiện để tìm task nhanh.')
add_number('Bấm nút thêm task nếu có quyền quản lý task.')
add_number('Nhập tiêu đề, mô tả, ưu tiên, ngày bắt đầu, hạn hoàn thành và người được giao.')
add_number('Bấm vào nút mũi tên để mở hoặc đóng cây task. Nút này chỉ dùng cho expand/collapse.')
add_number('Bấm vào tên task hoặc thanh timeline để mở panel chi tiết task.')
add_placeholder('P11', 'Ảnh chụp cây task trong tab Công việc, thấy rõ nút mũi tên expand/collapse và tên task có thể bấm.')
add_placeholder('P12', 'Ảnh chụp modal Thêm task mới.')

add_heading('6.5 Panel chi tiết task', 2)
add_number('Tab Tổng quan: xem mô tả, tiến độ, trạng thái, người thực hiện và thông tin quan trọng.')
add_number('Tab Subtask & Checklist: xem toàn bộ subtask, tạo subtask và theo dõi checklist.')
add_number('Tab Nhật ký: xem các cập nhật, log thao tác và trao đổi liên quan.')
add_number('Khi task đang mở, ưu tiên đọc Tổng quan trước rồi mới sang Checklist và Nhật ký.')
add_placeholder('P13', 'Ảnh chụp panel chi tiết task - tab Tổng quan, có tiến độ và trạng thái.')
add_placeholder('P14', 'Ảnh chụp panel chi tiết task - tab Subtask & Checklist.')
add_placeholder('P15', 'Ảnh chụp panel chi tiết task - tab Nhật ký.')

add_heading('6.6 Timeline / Gantt', 2)
add_number('Timeline hiển thị task theo mốc thời gian để nhìn nhanh task nào sắp đến hạn hoặc đang quá hạn.')
add_number('Bấm vào thanh timeline hoặc thanh task để mở task chi tiết.')
add_number('Nếu task chưa có ngày, nó sẽ nằm ở khu vực task chưa đặt mốc thời gian.')
add_placeholder('P16', 'Ảnh chụp khu Timeline / Gantt của dự án.')

add_heading('7. Ràng buộc nghiệp vụ cần nhớ', 1)
add_bullet('Subtask phải nằm trong khoảng ngày của task cha nếu muốn tạo subtask nhanh.')
add_bullet('Task đã hoàn thành thì không nên xem là có thể chỉnh sửa tự do; hệ thống có ràng buộc riêng cho trạng thái Done.')
add_bullet('Danh sách thành viên và vai trò phải đúng trước khi giao nhiều task.')
add_bullet('Nút mũi tên ở task chỉ dùng để mở/đóng cây task, không phải nút mở chi tiết.')
add_bullet('Muốn xem chi tiết task, hãy bấm vào tên task hoặc thanh timeline liên quan.')

add_heading('8. Quy trình làm việc đề xuất', 1)
add_number('Owner tạo dự án và mốc thời gian tổng.')
add_number('Owner thêm thành viên và phân vai trò ngay từ đầu.')
add_number('Manager/Staff đọc task được giao, mở chi tiết và cập nhật đúng phần việc của mình.')
add_number('Theo dõi timeline mỗi ngày để phát hiện task quá hạn.')
add_number('Dùng checklist để tránh quên bước khi task có nhiều bước nhỏ.')
add_number('Khi xong việc, kiểm tra nhật ký task để xác nhận trạng thái đã được ghi nhận.')

add_heading('9. Danh mục hình ảnh cần chèn', 1)
add_bullet('P01 - Danh sách dự án: chụp toàn bộ màn hình, thấy bộ lọc trạng thái và thẻ dự án.')
add_bullet('P02 - Công việc ở góc nhìn Manager: chụp cây task và timeline, có thể thấy task được quản lý.')
add_bullet('P03 - Panel chi tiết task cho Manager: chụp tab Tổng quan.')
add_bullet('P04 - Công việc ở góc nhìn Staff: chụp task có thể bấm vào để mở chi tiết.')
add_bullet('P05 - Panel chi tiết task cho Staff: chụp tab Subtask & Checklist.')
add_bullet('P06 - Danh sách dự án.')
add_bullet('P07 - Form Tạo dự án.')
add_bullet('P08 - Form Sửa dự án.')
add_bullet('P09 - Tab Thành viên.')
add_bullet('P10 - Popup Thêm thành viên.')
add_bullet('P11 - Cây task trong tab Công việc.')
add_bullet('P12 - Modal Thêm task mới.')
add_bullet('P13 - Panel chi tiết task - Tổng quan.')
add_bullet('P14 - Panel chi tiết task - Subtask & Checklist.')
add_bullet('P15 - Panel chi tiết task - Nhật ký.')
add_bullet('P16 - Timeline / Gantt.')

add_heading('10. Cách chèn ảnh vào tài liệu', 1)
add_bullet('Mở file DOCX bằng Microsoft Word.')
add_bullet('Tìm dòng [VI TRI CHEN ANH - Pxx].')
add_bullet('Chèn ảnh screenshot tương ứng ngay dưới dòng đó hoặc thay thế đoạn mô tả.')
add_bullet('Nếu muốn tài liệu gọn hơn, giữ lại nhãn Pxx và xoá phần mô tả sau khi chèn ảnh xong.')

add_heading('11. Lỗi thường gặp và cách xử lý', 1)
add_bullet('Không thấy nút sửa/lưu trữ: kiểm tra lại quyền của tài khoản.')
add_bullet('Không mở được chi tiết task: bấm vào tên task hoặc thanh timeline, không bấm vào nút mũi tên expand/collapse.')
add_bullet('Không tạo được subtask: kiểm tra task cha có đủ ngày bắt đầu và ngày kết thúc hay chưa.')
add_bullet('Không thấy toàn bộ dự án: có thể đang ở vai trò Staff/Viewer nên hệ thống giới hạn phạm vi hiển thị.')

add_heading('12. Gợi ý sử dụng tài liệu', 1)
add_bullet('Dùng bản này cho đào tạo Staff/Manager vì đã tách luồng sử dụng theo vai trò.')
add_bullet('Nếu cần bản cho Admin/Owner, có thể sao chép và bổ sung riêng các thao tác quản trị.')
add_bullet('Khi hệ thống thay đổi UI, chỉ cần cập nhật lại phần mô tả ảnh và các bước tương ứng.')

doc.save(output_path)
print(f'Generated: {output_path}')
