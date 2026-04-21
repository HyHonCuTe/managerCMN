from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import date

output_path = r"e:\Study\C#\managerCMN\docs\manual\HUONG_DAN_SU_DUNG_CHUC_NANG_DU_AN_CHI_TIET.docx"
prompt_path = r"e:\Study\C#\managerCMN\docs\manual\PROMPT_VIET_DOCX_HUONG_DAN_DU_AN.txt"

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

def add_title(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def add_h1(text):
    doc.add_heading(text, level=1)

def add_h2(text):
    doc.add_heading(text, level=2)

def add_bullet(text):
    doc.add_paragraph(text, style='List Bullet')

def add_num(text):
    doc.add_paragraph(text, style='List Number')

def add_image_placeholder(code, note):
    p = doc.add_paragraph()
    r = p.add_run(f"[VI TRI CHEN ANH - {code}]\\n{note}")
    r.bold = True

add_title('HUONG DAN SU DUNG CHUC NANG DU AN (PROJECT) - CHI TIET')
sub = doc.add_paragraph(f"He thong: managerCMN | Ngay cap nhat: {date.today().strftime('%d/%m/%Y')}")
sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph('Tai lieu nay huong dan day du cac thao tac trong module Du an, bao gom tao du an, quan ly thanh vien, quan ly task/subtask, cap nhat tien do, checklist, timeline va luu tru/xoa du an theo dung quyen. Tai lieu duoc viet de dung cho dao tao noi bo va huong dan handover.')

add_h1('1. Pham vi va doi tuong su dung')
add_bullet('Pham vi: Module Du an tai menu trai: Du an > Danh sach du an > Chi tiet du an.')
add_bullet('Doi tuong: Admin he thong, Project Owner, Project Manager, Staff, Viewer.')
add_bullet('Muc tieu: Thuc hien dung quy trinh, dung quyen, dung luong thao tac de tranh sai du lieu.')

add_h1('2. Truoc khi bat dau')
add_bullet('Tai khoan da dang nhap he thong managerCMN.')
add_bullet('Tai khoan co quyen truy cap module Du an.')
add_bullet('Neu huong dan kem hinh, chup anh man hinh theo danh muc "Vi tri chen anh" trong tai lieu nay.')
add_image_placeholder('P01', 'Menu trai voi muc Du an duoc to sang.')

add_h1('3. Tong quan vai tro va quyen')
add_h2('3.1 Nhom vai tro')
add_bullet('System Admin: Xem tat ca du an, co quyen quan tri cao nhat trong module Du an.')
add_bullet('Project Owner: Chu du an, co quyen quan ly du an/thanh vien/task.')
add_bullet('Project Manager: Quan ly task theo phan quyen du an (khong phai moi truong hop deu duoc sua metadata du an).')
add_bullet('Project Staff: Thuc hien cong viec duoc giao.')
add_bullet('Project Viewer: Chi xem, khong thao tac sua doi.')

add_h2('3.2 Quy tac quyen quan trong')
add_bullet('Tao/Sua thong tin du an: Owner va Admin.')
add_bullet('Luu tru du an: Owner va Admin.')
add_bullet('Xoa vinh vien du an: Chi Admin, va chi khi du an da luu tru.')
add_bullet('Them/xoa thanh vien: Owner va Admin (theo giao dien hien tai).')
add_bullet('Tao task, giao viec, cap nhat status/progress: Nhom co quyen quan ly task.')
add_bullet('Task da Done khong the rollback nguoc trang thai qua giao dien nhanh.')

add_h1('4. Quy trinh su dung theo man hinh')
add_h2('4.1 Man hinh Danh sach du an')
add_num('Vao menu Du an.')
add_num('Su dung bo loc trang thai de xem danh sach theo nhu cau: Tat ca, Planning, Active, On Hold, Completed, Cancelled, Archived.')
add_num('Bam vao the du an de vao Chi tiet.')
add_num('Bam nut "Tao du an" neu can tao moi.')
add_image_placeholder('P02', 'Man hinh danh sach du an + bo loc trang thai.')

add_h2('4.2 Tao du an moi')
add_num('Bam "Tao du an".')
add_num('Nhap: Ten du an (bat buoc), Mo ta, Ngay bat dau, Ngay ket thuc.')
add_num('Luu y: Ngay ket thuc phai sau ngay bat dau.')
add_num('Bam "Tao du an" de hoan tat.')
add_image_placeholder('P03', 'Form Tao du an moi.')

add_h2('4.3 Sua thong tin du an')
add_num('Tai man hinh Chi tiet du an, bam "Sua".')
add_num('Cap nhat thong tin: Ten, Mo ta, Trang thai, Start/End date.')
add_num('Bam "Luu".')
add_bullet('Neu du an da luu tru thi chi xem, khong sua.')
add_image_placeholder('P04', 'Form Sua du an.')

add_h2('4.4 Luu tru va xoa vinh vien du an')
add_num('Luu tru: Bam nut "Luu tru" tren trang Chi tiet du an.')
add_num('Sau khi luu tru, du an o che do chi xem.')
add_num('Xoa vinh vien: Chi Admin thay nut "Xoa vinh vien" tren du an da luu tru.')
add_bullet('Hanh dong xoa vinh vien khong the hoan tac.')
add_image_placeholder('P05', 'Nut Luu tru / Xoa vinh vien tren header du an.')

add_h2('4.5 Quan ly thanh vien du an')
add_num('Vao tab "Thanh vien" trong Chi tiet du an.')
add_num('Bam "Them thanh vien".')
add_num('Tim kiem theo ten/ma/phong ban, loc theo phong ban, tick nhieu nguoi neu can.')
add_num('Chon vai tro (Manager/Staff/Viewer) va bam "Them".')
add_num('Dung menu 3 cham de doi vai tro hoac xoa thanh vien khoi du an.')
add_image_placeholder('P06', 'Tab Thanh vien + popup Them thanh vien nhieu nguoi.')

add_h2('4.6 Quan ly task tren tab Cong viec')
add_num('Tai tab Cong viec, su dung bo loc Trang thai/Nguoi thuc hien neu can.')
add_num('Bam "Them task" de tao task goc.')
add_num('Nhap thong tin task: Tieu de, Mo ta, Task cha (neu la subtask), do uu tien, ngay bat dau/ket thuc, gio uoc tinh, nguoi duoc giao.')
add_num('Bam vao dong task de mo panel chi tiet task o ben phai.')
add_num('Nut mui ten o dau dong task dung de dong/mo cay task (expand/collapse).')
add_image_placeholder('P07', 'Tab Cong viec voi task tree va bo loc.')
add_image_placeholder('P08', 'Modal Them task moi.')

add_h2('4.7 Panel chi tiet task (3 tab)')
add_bullet('Tab Tong quan: Mo ta, tien do, trang thai, nguoi thuc hien, cap nhat phan cong.')
add_bullet('Tab Subtask & Checklist: danh sach subtask, tao subtask nhanh, checklist.')
add_bullet('Tab Nhat ky: cap nhat log, ghi chu, dinh kem tep (neu co cau hinh).')
add_num('Cap nhat tien do bang slider (task mode Manual).')
add_num('Danh dau hoan thanh nhanh bang nut "Danh dau hoan thanh" khi du dieu kien.')
add_num('Cap nhat trang thai qua dropdown Trang thai.')
add_num('Checklist da hoan thanh thi khong bo tick nguoc qua thao tac nhanh.')
add_image_placeholder('P09', 'Panel Chi tiet task - tab Tong quan.')
add_image_placeholder('P10', 'Panel Chi tiet task - tab Subtask & Checklist.')
add_image_placeholder('P11', 'Panel Chi tiet task - tab Nhat ky.')

add_h2('4.8 Quy tac tao Subtask trong panel task')
add_bullet('Task cha phai co du ngay bat dau va ngay ket thuc moi tao duoc subtask nhanh.')
add_bullet('Ngay cua subtask phai nam trong khoang ngay cua task cha.')
add_bullet('Neu task cha thieu moc thoi gian, he thong canh bao va chan tao subtask nhanh.')

add_h2('4.9 Timeline/Gantt')
add_num('Khu Timeline hien thi tat ca task/subtask tren truc thoi gian du an.')
add_num('Bam thanh timeline de mo nhanh chi tiet task.')
add_num('Nhom "Task chua dat moc thoi gian" giup theo doi cac task can bo sung ngay.')
add_image_placeholder('P12', 'Khu Timeline/Gantt va cac thanh task.')

add_h1('5. Luong thao tac de xuat (Best practice)')
add_num('Tao du an va dat moc thoi gian tong.')
add_num('Them thanh vien va gan dung vai tro truoc khi tao nhieu task.')
add_num('Tao task goc theo nhom cong viec lon, sau do tach thanh subtask co deadline ro rang.')
add_num('Gan nguoi phu trach ngay khi tao task de tranh task mo.')
add_num('Cap nhat tien do/ngat trang thai theo ngay hoac theo moc sprint.')
add_num('Dung checklist cho cac buoc lap lai de tranh bo sot.')
add_num('Kiem tra khu qua han va timeline moi ngay.')

add_h1('6. Danh muc hinh anh de chen vao tai lieu')
add_bullet('P01: Menu trai - muc Du an.')
add_bullet('P02: Danh sach du an + bo loc.')
add_bullet('P03: Form Tao du an.')
add_bullet('P04: Form Sua du an.')
add_bullet('P05: Nut Luu tru / Xoa vinh vien.')
add_bullet('P06: Popup Them thanh vien + tim kiem/loc/chon nhieu.')
add_bullet('P07: Task tree + bo loc task.')
add_bullet('P08: Modal Them task.')
add_bullet('P09-P11: 3 tab trong panel chi tiet task.')
add_bullet('P12: Timeline/Gantt.')

doc.add_paragraph('Meo chen anh nhanh trong Word: Click vao dong [VI TRI CHEN ANH - Pxx] -> Insert -> Pictures -> chon anh chup tu thu muc docs/images/...')

add_h1('7. Loi thuong gap va cach xu ly')
add_bullet('Khong thay du an: kiem tra vai tro va quyen thanh vien du an.')
add_bullet('Khong sua duoc du an: du an co the da luu tru hoac tai khoan khong phai Owner/Admin.')
add_bullet('Khong tao duoc subtask nhanh: task cha chua co day du StartDate va DueDate.')
add_bullet('Khong the rollback task da Done: day la rang buoc nghiep vu hien tai.')
add_bullet('Bam mui ten task ma mo panel chi tiet: can cap nhat file JS moi nhat va hard refresh trinh duyet.')

add_h1('8. Phu luc prompt de cap nhat tai lieu bang AI')
prompt_text = """Ban la Technical Writer cho he thong managerCMN.
Hay viet lai tai lieu huong dan su dung chi tiet cho module Du an (Project) duoi dang DOCX, tieng Viet, ro rang cho nguoi dung noi bo.

Yeu cau bat buoc:
1) Bao gom day du cac man hinh: Danh sach du an, Tao/Sua du an, Chi tiet du an, tab Thanh vien, tab Cong viec, panel chi tiet task (Tong quan, Subtask & Checklist, Nhat ky), Timeline/Gantt.
2) Trinh bay theo tung vai tro: Admin, Owner, Manager, Staff, Viewer; neu quyen khac nhau phai ghi ro.
3) Co huong dan thao tac theo buoc (step-by-step), co luong best practice de van hanh du an.
4) Co muc canh bao/rang buoc nghiep vu (VD: subtask phai nam trong khoang thoi gian task cha, task da Done khong rollback nhanh).
5) Chen cac placeholder anh theo mau: [VI TRI CHEN ANH - P01], [VI TRI CHEN ANH - P02]... kem chu thich can chup gi.
6) Van phong de doc, de train user moi; ngan gon nhung khong thieu chi tiet quan trong.
7) Ket thuc bang muc Loi thuong gap va cach xu ly.

Dau ra mong muon:
- 01 file DOCX dat ten: HUONG_DAN_SU_DUNG_CHUC_NANG_DU_AN_CHI_TIET.docx
- Co muc "Danh muc hinh anh can chen" de bo sung anh sau.
"""

doc.add_paragraph(prompt_text)

doc.save(output_path)
with open(prompt_path, 'w', encoding='utf-8') as f:
    f.write(prompt_text)

print(f"Generated: {output_path}")
print(f"Generated: {prompt_path}")
