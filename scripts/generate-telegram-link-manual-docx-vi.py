from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


OUTPUT_PATH = r"e:\Study\C#\managerCMN\docs\manual\HUONG_DAN_KET_NOI_TELEGRAM_CMN_MANAGEMENT.docx"


doc = Document()


def set_doc_defaults() -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        try:
            doc.styles[style_name].font.name = "Times New Roman"
        except Exception:
            pass


def add_paragraph(text: str = "", *, bold: bool = False, size: int | None = None, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.bold = bold
    if size is not None:
        r.font.size = Pt(size)
    return p


def add_heading(text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Times New Roman"
    return p


def add_bullet(text: str):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    return p


def add_number(text: str):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    return p


def add_placeholder(code: str, description: str):
    p = doc.add_paragraph()
    r = p.add_run(f"[CHÈN HÌNH {code}] {description}")
    r.bold = True
    r.italic = True
    r.font.name = "Times New Roman"
    return p


def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = header
        for p in hdr[idx].paragraphs:
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.bold = True

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            for p in cells[idx].paragraphs:
                for r in p.runs:
                    r.font.name = "Times New Roman"

    return table


def add_cover():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("HƯỚNG DẪN KẾT NỐI TELEGRAM\nVỚI CMN MANAGEMENT")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(20)

    add_paragraph("Tài liệu nội bộ dành cho người dùng công ty", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(f"Phiên bản: 1.0 | Ngày cập nhật: {date.today().strftime('%d/%m/%Y')}", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph("\n")
    add_placeholder("COVER", "Ảnh bìa hoặc logo CMN ở giữa trang đầu.")


def add_intro():
    add_heading("1. Giới thiệu", 1)
    add_paragraph(
        "Tài liệu này hướng dẫn cách kết nối tài khoản Telegram cá nhân với hệ thống CMN Management "
        "để nhận thông báo tự động khi có giao việc, duyệt đơn, cập nhật task hoặc các sự kiện hệ thống khác."
    )
    add_paragraph(
        "Nội dung được viết bằng tiếng Việt có dấu đầy đủ, trình bày theo từng bước, có chỗ chèn ảnh minh hoạ "
        "để thuận tiện khi chuyển thành tài liệu hướng dẫn chính thức."
    )


def add_prerequisites():
    add_heading("2. Điều kiện trước khi kết nối", 1)
    add_bullet("Người dùng đã có tài khoản đăng nhập CMN Management.")
    add_bullet("Đã cài đặt ứng dụng Telegram trên điện thoại hoặc máy tính.")
    add_bullet("Đang truy cập đúng tài khoản cá nhân cần liên kết.")
    add_bullet("Có quyền mở trang Hồ sơ cá nhân trong hệ thống.")
    add_placeholder("1", "Ảnh chụp màn hình menu Hồ sơ cá nhân trong hệ thống.")


def add_steps():
    add_heading("3. Các bước thực hiện", 1)

    add_heading("3.1 Mở trang Kết nối Telegram", 2)
    add_number("Đăng nhập vào CMN Management.")
    add_number("Mở menu Hồ sơ cá nhân.")
    add_number("Chọn mục Kết nối Telegram.")
    add_placeholder("2", "Ảnh chụp trang Kết nối Telegram với mã kết nối hiển thị rõ ràng.")

    add_heading("3.2 Sao chép mã kết nối", 2)
    add_number("Quan sát mã kết nối được hiển thị trên màn hình.")
    add_number("Mã này chỉ dùng một lần và sẽ tự hết hạn sau 15 phút.")
    add_number("Nếu quá thời gian, làm mới trang để lấy mã mới.")
    add_placeholder("3", "Ảnh cận cảnh mã kết nối và thời gian hết hạn 15 phút.")

    add_heading("3.3 Mở Telegram và tìm bot", 2)
    add_number("Mở Telegram.")
    add_number("Tìm bot theo username hiển thị trong hệ thống, ví dụ @cmn_notify_bot.")
    add_number("Mở khung chat với bot đó.")
    add_placeholder("4", "Ảnh chụp màn hình Telegram mở đúng bot cần kết nối.")

    add_heading("3.4 Gửi lệnh /start kèm mã kết nối", 2)
    add_number("Nhập lệnh /start kèm mã kết nối, ví dụ: /start ABCD1234.")
    add_number("Gửi tin nhắn tới bot.")
    add_number("Chờ bot phản hồi xác nhận liên kết thành công.")
    add_placeholder("5", "Ảnh chụp màn hình tin nhắn /start kèm mã kết nối đã gửi.")

    add_heading("3.5 Nhận thông báo thành công", 2)
    add_number("Nếu kết nối hợp lệ, bot sẽ phản hồi thông báo thành công.")
    add_number("Từ thời điểm đó, hệ thống sẽ gửi thông báo Telegram cho tài khoản đã liên kết.")
    add_placeholder("6", "Ảnh chụp màn hình bot phản hồi 'Kết nối thành công'.")


def add_check_and_errors():
    add_heading("4. Cách kiểm tra kết nối thành công", 1)
    add_bullet("Quay lại trang Kết nối Telegram trong hệ thống.")
    add_bullet("Xác nhận trạng thái đã kết nối.")
    add_bullet("Thử nhắn tin hoặc chờ một thông báo hệ thống để kiểm tra việc nhận thông báo qua Telegram.")
    add_placeholder("7", "Ảnh chụp màn hình trạng thái đã kết nối thành công trong hồ sơ cá nhân.")

    add_heading("5. Các lỗi thường gặp và cách xử lý", 1)
    add_table(
        ["Triệu chứng", "Nguyên nhân", "Cách xử lý"],
        [
            ["Bot không phản hồi", "Chưa gửi đúng /start kèm mã hoặc mã đã hết hạn.", "Làm mới trang để lấy mã mới và gửi lại /start."],
            ["Mã không hợp lệ", "Gửi nhầm bot hoặc sai mã kết nối.", "Kiểm tra đúng bot và nhập lại mã chính xác."],
            ["Không nhận thông báo", "Tài khoản chưa liên kết hoặc bot chưa được cấp quyền gửi.", "Kiểm tra trạng thái liên kết và thử lại sau khi kết nối thành công."],
            ["Mã hết hạn", "Quá 15 phút kể từ lúc tạo mã.", "Làm mới trang Kết nối Telegram để lấy mã mới."],
        ],
    )


def add_security_notes():
    add_heading("6. Lưu ý bảo mật", 1)
    add_bullet("Mã kết nối chỉ dùng một lần.")
    add_bullet("Mã sẽ tự hết hạn sau 15 phút.")
    add_bullet("Không chia sẻ mã cho người khác.")
    add_bullet("Nếu nghi ngờ mã bị lộ, hãy làm mới trang để tạo mã mới ngay.")
    add_bullet("Chỉ kết nối Telegram trên thiết bị cá nhân an toàn của bạn.")
    add_placeholder("8", "Ảnh chụp cảnh báo bảo mật hoặc phần mô tả mã kết nối hết hạn 15 phút.")


def add_conclusion():
    add_heading("7. Kết luận", 1)
    add_paragraph(
        "Sau khi hoàn tất kết nối, người dùng sẽ nhận thông báo nhanh hơn và không cần đăng nhập lại nhiều lần để kiểm tra. "
        "Khi cần liên kết lại, chỉ cần tạo mã mới từ trang Kết nối Telegram và thực hiện lại các bước ở trên."
    )


def main():
    set_doc_defaults()
    add_cover()

    doc.add_page_break()
    add_paragraph("Mục lục", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph("1. Giới thiệu")
    add_paragraph("2. Điều kiện trước khi kết nối")
    add_paragraph("3. Các bước thực hiện")
    add_paragraph("4. Cách kiểm tra kết nối thành công")
    add_paragraph("5. Các lỗi thường gặp và cách xử lý")
    add_paragraph("6. Lưu ý bảo mật")
    add_paragraph("7. Kết luận")

    add_intro()
    add_prerequisites()
    add_steps()
    add_check_and_errors()
    add_security_notes()
    add_conclusion()

    doc.save(OUTPUT_PATH)
    print(f"Generated: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()