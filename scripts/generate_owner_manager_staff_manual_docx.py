from datetime import date
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

OUTPUT_PATH = r"e:\\Study\\C#\\managerCMN\\docs\\manual\\HUONG_DAN_SU_DUNG_DU_AN_OWNER_MANAGER_STAFF_CHI_TIET.docx"


doc = Document()

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)


def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(18)


def add_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = p.add_run(text)
    r.font.name = 'Times New Roman'


def add_h1(text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.name = 'Times New Roman'


def add_h2(text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.name = 'Times New Roman'


def add_paragraph(text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.name = 'Times New Roman'


def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    r.font.name = 'Times New Roman'


def add_number(text):
    p = doc.add_paragraph(style='List Number')
    r = p.add_run(text)
    r.font.name = 'Times New Roman'


def add_placeholder(code, desc):
    p = doc.add_paragraph()
    r = p.add_run(f"[VI TRI CHEN ANH - {code}]\\n{desc}")
    r.bold = True
    r.font.name = 'Times New Roman'


def add_permission_table():
    headers = ["Chuc nang", "Owner", "Manager", "Staff", "Ghi chu"]
    rows = [
        ["Xem danh sach du an", "Co", "Co", "Co", "Theo quyen thanh vien du an"],
        ["Tao du an", "Co", "Khong", "Khong", "Admin he thong cung co"],
        ["Sua thong tin du an", "Co", "Khong", "Khong", "Gom ten, mo ta, moc ngay"],
        ["Luu tru du an", "Co", "Khong", "Khong", "Admin he thong cung co"],
        ["Phuc hoi du an tu luu tru", "Khong", "Khong", "Khong", "Chi Admin he thong"],
        ["Them thanh vien vao du an", "Co", "Co", "Khong", "Manager duoc them staff theo quy trinh moi"],
        ["Doi vai tro thanh vien", "Co", "Han che", "Khong", "Manager khong nang cap Owner/Manager"],
        ["Tao task/task cha", "Co", "Theo quyen", "Khong", "Manager theo pham vi task duoc giao"],
        ["Gan nguoi thuc hien task", "Co", "Co", "Khong", "Manager gan task trong pham vi duoc phep"],
        ["Cap nhat progress/status task", "Co", "Co", "Co", "Staff cap nhat task duoc giao"],
        ["Quan ly checklist", "Co", "Co", "Co", "Theo quyen task"],
        ["Xem timeline/Gantt", "Co", "Co", "Co", "Du lieu theo pham vi quyen"],
    ]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    header_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        header_cells[i].text = h
        for p in header_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = 'Times New Roman'

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.name = 'Times New Roman'

    for row in table.rows:
        row.cells[0].width = Cm(5.0)
        row.cells[1].width = Cm(2.1)
        row.cells[2].width = Cm(2.1)
        row.cells[3].width = Cm(2.1)
        row.cells[4].width = Cm(6.0)


add_title("HUONG DAN SU DUNG MODULE DU AN")
add_title("BAN DAY DU CHO OWNER / MANAGER / STAFF")
add_subtitle(f"He thong: managerCMN | Ngay cap nhat: {date.today().strftime('%d/%m/%Y')}")
add_paragraph("Tai lieu nay tong hop day du huong dan su dung module Du an cho 3 vai tro Owner, Manager, Staff. Noi dung theo phong cach thao tac thuc te, co vi tri chen anh de bo sung screenshot noi bo.")

add_h1("1. Muc tieu va pham vi")
add_bullet("Muc tieu: Giup user moi thao tac dung quyen, dung luong va dung man hinh trong module Du an.")
add_bullet("Pham vi: Danh sach du an, chi tiet du an, task/subtask/checklist, timeline, thanh vien, luu tru/phuc hoi.")
add_bullet("Doi tuong: Project Owner, Project Manager, Project Staff.")
add_placeholder("OMGS-01", "Man hinh menu ben trai, muc Du an.")

add_h1("2. Tong quan vai tro va quyen")
add_paragraph("Bang duoi day la ma tran quyen de user nhanh xac dinh viec minh co duoc phep thao tac hay khong.")
add_permission_table()
add_placeholder("OMGS-02", "Anh man hinh chi tiet du an voi role badge (Owner/Manager/Staff).")

add_h1("3. Luong tong quan cua module")
add_number("Vao Danh sach du an -> chon du an can thao tac.")
add_number("Mo Chi tiet du an -> xac dinh tab can dung: Cong viec, Thanh vien, Timeline.")
add_number("Mo task chi tiet -> cap nhat trang thai/progress/checklist/nhat ky.")
add_number("Theo doi qua han va tien do tong de xu ly som cac nut that.")
add_placeholder("OMGS-03", "Danh sach du an: bo loc trang thai + card KPI.")

add_h1("4. Huong dan chi tiet cho OWNER")
add_h2("4.1 Tao du an moi")
add_number("Tai trang Danh sach du an, bam Tao du an.")
add_number("Nhap ten, mo ta, ngay bat dau, ngay ket thuc.")
add_number("Luu du an va kiem tra du an da hien trong danh sach.")
add_placeholder("OMGS-04", "Form Tao du an.")

add_h2("4.2 Chinh sua du an")
add_number("Mo chi tiet du an -> bam Sua.")
add_number("Cap nhat thong tin tong quan va trang thai phu hop.")
add_number("Luu thay doi -> kiem tra thong bao thanh cong.")
add_placeholder("OMGS-05", "Form Sua du an + thong bao cap nhat thanh cong.")

add_h2("4.3 Quan ly thanh vien")
add_number("Vao tab Thanh vien.")
add_number("Bam Them thanh vien.")
add_number("Tim theo ten/ma nhan vien/phong ban.")
add_number("Chon role phu hop va xac nhan them.")
add_number("Co the doi role hoac xoa thanh vien bang menu hanh dong.")
add_placeholder("OMGS-06", "Popup Them thanh vien, tim kiem + loc phong ban.")

add_h2("4.4 Tao cau truc task")
add_number("Vao tab Cong viec -> Them task cha cho tung nhom cong viec lon.")
add_number("Mo task cha -> tao subtask theo tung dau muc con.")
add_number("Gan assignee va due date ro rang cho tung task.")
add_number("Kiem tra timeline de dam bao khong xung dot tien do.")
add_placeholder("OMGS-07", "Task tree voi task cha/subtask.")

add_h2("4.5 Luu tru du an")
add_number("Khi du an da ket thuc, Owner co the bam Luu tru.")
add_number("Sau luu tru, du an sang che do read-only.")
add_number("Neu can phuc hoi, can Admin he thong thao tac.")
add_placeholder("OMGS-08", "Nut Luu tru tai header Chi tiet du an.")

add_h1("5. Huong dan chi tiet cho MANAGER")
add_h2("5.1 Xem du an va KPI theo pham vi")
add_number("Vao du an duoc tham gia voi role Manager.")
add_number("Tai chi tiet du an, theo doi KPI trong pham vi task Manager duoc giao/quan ly.")
add_number("Uu tien xu ly task qua han va task gan den han.")
add_placeholder("OMGS-09", "KPI o role Manager (task trong pham vi cua Manager).")

add_h2("5.2 Them staff moi vao du an")
add_number("Vao tab Thanh vien.")
add_number("Bam Them thanh vien -> chon role Staff.")
add_number("Xac nhan them, kiem tra staff moi xuat hien trong danh sach.")
add_bullet("Luu y: Manager khong duoc bo nhiem Owner/Manager moi.")
add_placeholder("OMGS-10", "Manager them staff moi trong popup thanh vien.")

add_h2("5.3 Gan task cho staff moi")
add_number("Mo task trong pham vi Manager duoc phep quan ly.")
add_number("Tai panel chi tiet task, cap nhat danh sach nguoi thuc hien.")
add_number("Chon staff vua them va luu thay doi.")
add_number("Kiem tra nhat ky task de xac nhan hanh dong phan cong.")
add_placeholder("OMGS-11", "Panel task voi muc gan nguoi thuc hien cho staff moi.")

add_h2("5.4 Dieu phoi task hang ngay")
add_number("Dung bo loc theo nguoi thuc hien va trang thai de rap van de.")
add_number("Theo doi Timeline/Gantt de phat hien task tre han.")
add_number("Cap nhat ghi chu va huong dan xu ly trong Nhat ky task.")
add_placeholder("OMGS-12", "Timeline/Gantt + bo loc task theo assignee.")

add_h1("6. Huong dan chi tiet cho STAFF")
add_h2("6.1 Nhan viec")
add_number("Mo du an -> vao tab Cong viec.")
add_number("Tim cac task co ten minh trong danh sach nguoi thuc hien.")
add_number("Mo panel task de doc mo ta, due date, yeu cau dau ra.")
add_placeholder("OMGS-13", "Staff mo task duoc giao trong task tree.")

add_h2("6.2 Cap nhat progress va trang thai")
add_number("Tai tab Tong quan cua task, cap nhat progress theo muc do hoan thanh.")
add_number("Doi trang thai task dung voi thuc te cong viec.")
add_number("Khi dat tieu chi, danh dau hoan thanh theo quy trinh.")
add_placeholder("OMGS-14", "Staff cap nhat progress/status task.")

add_h2("6.3 Lam viec voi Subtask va Checklist")
add_number("Sang tab Subtask & Checklist.")
add_number("Danh dau checklist da xong theo tung buoc.")
add_number("Neu can, them ghi chu vao Nhat ky de bao cao blocker.")
add_placeholder("OMGS-15", "Staff thao tac checklist va subtask.")

add_h2("6.4 Nhat ky cong viec")
add_number("Sang tab Nhat ky.")
add_number("Ghi cap nhat ngan gon: da lam gi, vuong gi, can ai ho tro.")
add_number("Dinh kem tep (neu duoc bat) de lam bang chung dau viec.")
add_placeholder("OMGS-16", "Tab Nhat ky voi danh sach update moi nhat.")

add_h1("7. Huong dan theo man hinh chuc nang")
add_h2("7.1 Danh sach du an")
add_bullet("Bo loc theo status: Planning, Active, On Hold, Completed, Cancelled, Archived.")
add_bullet("Moi card hien tien do, so task, so thanh vien, canh bao qua han.")
add_placeholder("OMGS-17", "Trang Danh sach du an voi bo loc status.")

add_h2("7.2 Chi tiet du an")
add_bullet("KPI, thong tin tong quan, trang thai, ngay bat dau/ket thuc.")
add_bullet("Tab Cong viec: task tree + thao tac task.")
add_bullet("Tab Thanh vien: them/sua/xoa thanh vien theo quyen.")
add_placeholder("OMGS-18", "Trang Chi tiet du an day du cac khu vuc.")

add_h2("7.3 Task Tree va Timeline")
add_bullet("Mui ten o dau dong task: chi de expand/collapse cay task.")
add_bullet("Bam ten task hoac thanh timeline de mo chi tiet task.")
add_bullet("Task khong co moc ngay se nam trong nhom chua dat moc.")
add_placeholder("OMGS-19", "Task tree + timeline tren cung mot man hinh.")

add_h1("8. Rang buoc nghiep vu quan trong")
add_bullet("Subtask nhanh can task cha co day du StartDate va DueDate.")
add_bullet("Ngay subtask phai nam trong khoang cua task cha.")
add_bullet("Task da Done co rang buoc cap nhat trang thai theo luong cho phep.")
add_bullet("Project da Archived thi khong cho sua nhu du an dang hoat dong.")
add_bullet("Restore tu Archived hien la quyen cua Admin he thong.")

add_h1("9. Loi thuong gap va cach xu ly nhanh")
add_bullet("Khong thay nut hanh dong: kiem tra role hien tai va trang thai Archived.")
add_bullet("Khong them duoc subtask: bo sung ngay bat dau/ket thuc cho task cha.")
add_bullet("Khong thay task can lam: dung bo loc va kiem tra assignee cua task.")
add_bullet("Khong phan cong duoc staff: xac nhan staff da la thanh vien du an.")
add_bullet("KPI khong giong Owner: role Manager hien KPI theo pham vi task duoc giao.")

add_h1("10. Danh muc hinh anh can chen")
for i in range(1, 20):
    add_bullet(f"OMGS-{i:02d}: Chen anh tuong ung placeholder OMGS-{i:02d}.")

add_h1("11. Checklist onboarding 30 phut")
add_number("05 phut: Doc muc ma tran quyen Owner/Manager/Staff.")
add_number("10 phut: Tu tay mo 1 du an, vao task tree, mo panel task.")
add_number("10 phut: Thu cap nhat progress/checklist/nhat ky tren task demo.")
add_number("05 phut: On lai cac rang buoc nghiep vu va loi thuong gap.")

add_h1("12. Prompt mau de cap nhat tai lieu bang AI")
add_paragraph("Ban co the dung prompt trong file PROMPT_VIET_LAI_HDSD_OWNER_MANAGER_STAFF.txt de tai sinh tai lieu khi UI thay doi.")


doc.save(OUTPUT_PATH)
print(f"Generated: {OUTPUT_PATH}")
