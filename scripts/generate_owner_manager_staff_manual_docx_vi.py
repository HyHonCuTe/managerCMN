from datetime import date
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

OUTPUT_PATH = r"e:\\Study\\C#\\managerCMN\\docs\\manual\\HUONG_DAN_SU_DUNG_DU_AN_OWNER_MANAGER_STAFF_CHI_TIET_TIENG_VIET_CO_DAU.docx"


doc = Document()

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)


def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(18)


def add_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = p.add_run(text)
    r.font.name = 'Times New Roman'


def add_h1(text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.name = 'Times New Roman'


def add_h2(text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.name = 'Times New Roman'


def add_paragraph(text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.name = 'Times New Roman'


def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    r.font.name = 'Times New Roman'


def add_number(text):
    p = doc.add_paragraph(style='List Number')
    r = p.add_run(text)
    r.font.name = 'Times New Roman'


def add_placeholder(code, desc):
    p = doc.add_paragraph()
    r = p.add_run(f"[VI TRI CHEN ANH - {code}]\\n{desc}")
    r.bold = True
    r.font.name = 'Times New Roman'


def add_permission_table():
    headers = ["Chức năng", "Owner", "Manager", "Staff", "Ghi chú"]
    rows = [
        ["Xem danh sách dự án", "Có", "Có", "Có", "Theo quyền thành viên dự án"],
        ["Tạo dự án", "Có", "Không", "Không", "Admin hệ thống cũng có"],
        ["Sửa thông tin dự án", "Có", "Không", "Không", "Gồm tên, mô tả, mốc ngày"],
        ["Lưu trữ dự án", "Có", "Không", "Không", "Admin hệ thống cũng có"],
        ["Phục hồi dự án từ lưu trữ", "Không", "Không", "Không", "Chỉ Admin hệ thống"],
        ["Thêm thành viên vào dự án", "Có", "Có", "Không", "Manager được thêm staff theo quy trình mới"],
        ["Đổi vai trò thành viên", "Có", "Hạn chế", "Không", "Manager không nâng cấp Owner/Manager"],
        ["Tạo task/task cha", "Có", "Theo quyền", "Không", "Manager theo phạm vi task được giao"],
        ["Gán người thực hiện task", "Có", "Có", "Không", "Manager gán task trong phạm vi được phép"],
        ["Cập nhật progress/status task", "Có", "Có", "Có", "Staff cập nhật task được giao"],
        ["Quản lý checklist", "Có", "Có", "Có", "Theo quyền task"],
        ["Xem timeline/Gantt", "Có", "Có", "Có", "Dữ liệu theo phạm vi quyền"],
    ]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    header_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        header_cells[i].text = h
        for p in header_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = 'Times New Roman'

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.name = 'Times New Roman'

    for row in table.rows:
        row.cells[0].width = Cm(5.0)
        row.cells[1].width = Cm(2.1)
        row.cells[2].width = Cm(2.1)
        row.cells[3].width = Cm(2.1)
        row.cells[4].width = Cm(6.0)


add_title("HƯỚNG DẪN SỬ DỤNG MODULE DỰ ÁN")
add_title("BẢN ĐẦY ĐỦ CHO OWNER / MANAGER / STAFF")
add_subtitle(f"Hệ thống: managerCMN | Ngày cập nhật: {date.today().strftime('%d/%m/%Y')}")
add_paragraph("Tài liệu này tổng hợp đầy đủ hướng dẫn sử dụng module Dự án cho 3 vai trò Owner, Manager, Staff. Nội dung theo phong cách thao tác thực tế, có vị trí chèn ảnh để bổ sung screenshot nội bộ.")

add_h1("1. Mục tiêu và phạm vi")
add_bullet("Mục tiêu: Giúp user mới thao tác đúng quyền, đúng luồng và đúng màn hình trong module Dự án.")
add_bullet("Phạm vi: Danh sách dự án, chi tiết dự án, task/subtask/checklist, timeline, thành viên, lưu trữ/phục hồi.")
add_bullet("Đối tượng: Project Owner, Project Manager, Project Staff.")
add_placeholder("OMGS-01", "Màn hình menu bên trái, mục Dự án.")

add_h1("2. Tổng quan vai trò và quyền")
add_paragraph("Bảng dưới đây là ma trận quyền để user nhanh xác định việc mình có được phép thao tác hay không.")
add_permission_table()
add_placeholder("OMGS-02", "Ảnh màn hình chi tiết dự án với role badge (Owner/Manager/Staff).")

add_h1("3. Luồng tổng quan của module")
add_number("Vào Danh sách dự án -> chọn dự án cần thao tác.")
add_number("Mở Chi tiết dự án -> xác định tab cần dùng: Công việc, Thành viên, Timeline.")
add_number("Mở task chi tiết -> cập nhật trạng thái/progress/checklist/nhật ký.")
add_number("Theo dõi quá hạn và tiến độ tổng để xử lý sớm các nút thắt.")
add_placeholder("OMGS-03", "Danh sách dự án: bộ lọc trạng thái + card KPI.")

add_h1("4. Hướng dẫn chi tiết cho OWNER")
add_h2("4.1 Tạo dự án mới")
add_number("Tại trang Danh sách dự án, bấm Tạo dự án.")
add_number("Nhập tên, mô tả, ngày bắt đầu, ngày kết thúc.")
add_number("Lưu dự án và kiểm tra dự án đã hiện trong danh sách.")
add_placeholder("OMGS-04", "Form Tạo dự án.")

add_h2("4.2 Chỉnh sửa dự án")
add_number("Mở chi tiết dự án -> bấm Sửa.")
add_number("Cập nhật thông tin tổng quan và trạng thái phù hợp.")
add_number("Lưu thay đổi -> kiểm tra thông báo thành công.")
add_placeholder("OMGS-05", "Form Sửa dự án + thông báo cập nhật thành công.")

add_h2("4.3 Quản lý thành viên")
add_number("Vào tab Thành viên.")
add_number("Bấm Thêm thành viên.")
add_number("Tìm theo tên/mã nhân viên/phòng ban.")
add_number("Chọn role phù hợp và xác nhận thêm.")
add_number("Có thể đổi role hoặc xóa thành viên bằng menu hành động.")
add_placeholder("OMGS-06", "Popup Thêm thành viên, tìm kiếm + lọc phòng ban.")

add_h2("4.4 Tạo cấu trúc task")
add_number("Vào tab Công việc -> Thêm task cha cho từng nhóm công việc lớn.")
add_number("Mở task cha -> tạo subtask theo từng đầu mục con.")
add_number("Gán assignee và due date rõ ràng cho từng task.")
add_number("Kiểm tra timeline để đảm bảo không xung đột tiến độ.")
add_placeholder("OMGS-07", "Task tree với task cha/subtask.")

add_h2("4.5 Lưu trữ dự án")
add_number("Khi dự án đã kết thúc, Owner có thể bấm Lưu trữ.")
add_number("Sau lưu trữ, dự án sang chế độ read-only.")
add_number("Nếu cần phục hồi, cần Admin hệ thống thao tác.")
add_placeholder("OMGS-08", "Nút Lưu trữ tại header Chi tiết dự án.")

add_h1("5. Hướng dẫn chi tiết cho MANAGER")
add_h2("5.1 Xem dự án và KPI theo phạm vi")
add_number("Vào dự án được tham gia với role Manager.")
add_number("Tại chi tiết dự án, theo dõi KPI trong phạm vi task Manager được giao/quản lý.")
add_number("Ưu tiên xử lý task quá hạn và task gần đến hạn.")
add_placeholder("OMGS-09", "KPI ở role Manager (task trong phạm vi của Manager).")

add_h2("5.2 Thêm staff mới vào dự án")
add_number("Vào tab Thành viên.")
add_number("Bấm Thêm thành viên -> chọn role Staff.")
add_number("Xác nhận thêm, kiểm tra staff mới xuất hiện trong danh sách.")
add_bullet("Lưu ý: Manager không được bổ nhiệm Owner/Manager mới.")
add_placeholder("OMGS-10", "Manager thêm staff mới trong popup thành viên.")

add_h2("5.3 Gán task cho staff mới")
add_number("Mở task trong phạm vi Manager được phép quản lý.")
add_number("Tại panel chi tiết task, cập nhật danh sách người thực hiện.")
add_number("Chọn staff vừa thêm và lưu thay đổi.")
add_number("Kiểm tra nhật ký task để xác nhận hành động phân công.")
add_placeholder("OMGS-11", "Panel task với mục gán người thực hiện cho staff mới.")

add_h2("5.4 Điều phối task hằng ngày")
add_number("Dùng bộ lọc theo người thực hiện và trạng thái để ráp vấn đề.")
add_number("Theo dõi Timeline/Gantt để phát hiện task trễ hạn.")
add_number("Cập nhật ghi chú và hướng dẫn xử lý trong Nhật ký task.")
add_placeholder("OMGS-12", "Timeline/Gantt + bộ lọc task theo assignee.")

add_h1("6. Hướng dẫn chi tiết cho STAFF")
add_h2("6.1 Nhận việc")
add_number("Mở dự án -> vào tab Công việc.")
add_number("Tìm các task có tên mình trong danh sách người thực hiện.")
add_number("Mở panel task để đọc mô tả, due date, yêu cầu đầu ra.")
add_placeholder("OMGS-13", "Staff mở task được giao trong task tree.")

add_h2("6.2 Cập nhật progress và trạng thái")
add_number("Tại tab Tổng quan của task, cập nhật progress theo mức độ hoàn thành.")
add_number("Đổi trạng thái task đúng với thực tế công việc.")
add_number("Khi đạt tiêu chí, đánh dấu hoàn thành theo quy trình.")
add_placeholder("OMGS-14", "Staff cập nhật progress/status task.")

add_h2("6.3 Làm việc với Subtask và Checklist")
add_number("Sang tab Subtask & Checklist.")
add_number("Đánh dấu checklist đã xong theo từng bước.")
add_number("Nếu cần, thêm ghi chú vào Nhật ký để báo cáo blocker.")
add_placeholder("OMGS-15", "Staff thao tác checklist và subtask.")

add_h2("6.4 Nhật ký công việc")
add_number("Sang tab Nhật ký.")
add_number("Ghi cập nhật ngắn gọn: đã làm gì, vướng gì, cần ai hỗ trợ.")
add_number("Đính kèm tệp (nếu được bật) để làm bằng chứng đầu việc.")
add_placeholder("OMGS-16", "Tab Nhật ký với danh sách update mới nhất.")

add_h1("7. Hướng dẫn theo màn hình chức năng")
add_h2("7.1 Danh sách dự án")
add_bullet("Bộ lọc theo status: Planning, Active, On Hold, Completed, Cancelled, Archived.")
add_bullet("Mỗi card hiện tiến độ, số task, số thành viên, cảnh báo quá hạn.")
add_placeholder("OMGS-17", "Trang Danh sách dự án với bộ lọc status.")

add_h2("7.2 Chi tiết dự án")
add_bullet("KPI, thông tin tổng quan, trạng thái, ngày bắt đầu/kết thúc.")
add_bullet("Tab Công việc: task tree + thao tác task.")
add_bullet("Tab Thành viên: thêm/sửa/xóa thành viên theo quyền.")
add_placeholder("OMGS-18", "Trang Chi tiết dự án đầy đủ các khu vực.")

add_h2("7.3 Task Tree và Timeline")
add_bullet("Mũi tên ở đầu dòng task: chỉ để expand/collapse cây task.")
add_bullet("Bấm tên task hoặc thanh timeline để mở chi tiết task.")
add_bullet("Task không có mốc ngày sẽ nằm trong nhóm chưa đặt mốc.")
add_placeholder("OMGS-19", "Task tree + timeline trên cùng một màn hình.")

add_h1("8. Ràng buộc nghiệp vụ quan trọng")
add_bullet("Subtask nhanh cần task cha có đầy đủ StartDate và DueDate.")
add_bullet("Ngày subtask phải nằm trong khoảng của task cha.")
add_bullet("Task đã Done có ràng buộc cập nhật trạng thái theo luồng cho phép.")
add_bullet("Project đã Archived thì không cho sửa như dự án đang hoạt động.")
add_bullet("Restore từ Archived hiện là quyền của Admin hệ thống.")

add_h1("9. Lỗi thường gặp và cách xử lý nhanh")
add_bullet("Không thấy nút hành động: kiểm tra role hiện tại và trạng thái Archived.")
add_bullet("Không thêm được subtask: bổ sung ngày bắt đầu/kết thúc cho task cha.")
add_bullet("Không thấy task cần làm: dùng bộ lọc và kiểm tra assignee của task.")
add_bullet("Không phân công được staff: xác nhận staff đã là thành viên dự án.")
add_bullet("KPI không giống Owner: role Manager hiển thị KPI theo phạm vi task được giao.")

add_h1("10. Danh mục hình ảnh cần chèn")
for i in range(1, 20):
    add_bullet(f"OMGS-{i:02d}: Chèn ảnh tương ứng placeholder OMGS-{i:02d}.")

add_h1("11. Checklist onboarding 30 phút")
add_number("05 phút: Đọc mục ma trận quyền Owner/Manager/Staff.")
add_number("10 phút: Tự tay mở 1 dự án, vào task tree, mở panel task.")
add_number("10 phút: Thử cập nhật progress/checklist/nhật ký trên task demo.")
add_number("05 phút: Ôn lại các ràng buộc nghiệp vụ và lỗi thường gặp.")

add_h1("12. Prompt mẫu để cập nhật tài liệu bằng AI")
add_paragraph("Bạn có thể dùng prompt trong file PROMPT_VIET_LAI_HDSD_OWNER_MANAGER_STAFF.txt để tái sinh tài liệu khi UI thay đổi.")


doc.save(OUTPUT_PATH)
print(f"Generated: {OUTPUT_PATH}")
